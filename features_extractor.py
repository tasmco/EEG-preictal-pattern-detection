import os
import mne
import numpy as np
from docx import Document
from docx.shared import Pt

# دالة لحساب Sample Entropy
def sample_entropy(U, m=2, r=None):
    U = np.asarray(U)
    N = len(U)
    if r is None:
        r = 0.2 * np.std(U)
    def _phi(m):
        x = np.array([U[i: i + m] for i in range(N - m + 1)])
        C = 0
        for i in range(len(x)):
            dist = np.max(np.abs(x[i] - x), axis=1)
            C += np.sum(dist <= r) - 1  # استبعاد الحالة الذاتية
        return C / ((N - m) * (N - m + 1))
    phi_m = _phi(m)
    phi_m1 = _phi(m + 1)
    if phi_m == 0 or phi_m1 == 0:
        return np.nan
    return -np.log(phi_m1 / phi_m)

# دالة استخراج الميزات من سجل EEG (لكل قناة)
def extract_features(raw):
    features = {}
    data, times = raw[:]
    ch_names = raw.ch_names

    # تعريف نطاقات التردد (بالهرتز)
    bands = {
        'delta': (0.5, 4),
        'theta': (4, 8),
        'alpha': (8, 13),
        'beta':  (13, 30),
        'gamma': (30, 50)
    }
    for idx, ch in enumerate(ch_names):
        channel_data = data[idx]
        feats = {}
        # الميزات الزمنية
        feats[f'{ch}_mean'] = np.mean(channel_data)
        feats[f'{ch}_std'] = np.std(channel_data)
        
        # حساب PSD باستخدام الدالة compute_psd من كائن raw
        psd_result = raw.compute_psd(method='welch', picks=[idx], verbose=False, average=False)
        if isinstance(psd_result, tuple):
            psds, freqs = psd_result
        else:
            psds = psd_result.data
            freqs = psd_result.freqs

        # نأخذ أول قسم من النتائج (في حال وجود أكثر من قسم)
        psd = psds[0]
        feats[f'{ch}_total_power'] = np.sum(psd)
        for band, (fmin, fmax) in bands.items():
            idx_band = np.logical_and(freqs >= fmin, freqs < fmax)
            band_power = np.sum(psd[idx_band])
            feats[f'{ch}_{band}_power'] = band_power

        # الميزات غير الخطية: حساب Sample Entropy
        sampen = sample_entropy(channel_data, m=2, r=0.2 * np.std(channel_data))
        feats[f'{ch}_sampen'] = sampen

        features[ch] = feats
    return features

def main():
    # مسار الملفات المقسمة (بعد التطبيع وخمس ثواني لكل جزء)
    input_dir = r"C:\Users\XXX\five_second_segments"
    # مسار حفظ ملف استخراج الميزات النهائي
    output_file = r"C:\Users\XXX\features_extraction.docx"
    
    doc = Document()
    doc.add_heading("Feature Extraction Report", 0)
    
    # قراءة الملفات التي تنتهي بـ _raw.fif
    files = [f for f in os.listdir(input_dir) if f.endswith('_raw.fif')]
    
    for fname in files:
        file_path = os.path.join(input_dir, fname)
        print(f"Processing: {fname}")
        try:
            raw = mne.io.read_raw_fif(file_path, preload=True, verbose=False)
        except Exception as e:
            print(f"Error reading {fname}: {e}")
            continue
        
        feats = extract_features(raw)
        # إضافة عنوان الملف إلى التقرير
        doc.add_heading(fname, level=1)
        # لكل قناة، نضيف جدول الميزات
        for ch, feat_dict in feats.items():
            doc.add_heading(ch, level=2)
            table = doc.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Feature"
            hdr_cells[1].text = "Value"
            for key, value in feat_dict.items():
                row_cells = table.add_row().cells
                # إزالة اسم القناة من المفتاح لعرض اسم الميزة فقط
                row_cells[0].text = key.replace(f"{ch}_", "")
                if isinstance(value, float) and not np.isnan(value):
                    row_cells[1].text = f"{value:.6f}"
                else:
                    row_cells[1].text = str(value)
            doc.add_paragraph()  # فاصل بين الجداول
    
    doc.save(output_file)
    print(f"Features saved in: {output_file}")

if __name__ == "__main__":
    main()
