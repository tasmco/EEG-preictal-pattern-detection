import os
import glob
import mne
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# مسارات الإدخال والإخراج
input_dir = r"C:\Users\XXX\normalized_segments"
output_dir = r"C:\Users\XXX\normalizition_quality"

# إنشاء مجلد الإخراج إذا لم يكن موجودًا
if not os.path.exists(output_dir):
    os.makedirs(output_dir)

# مسار ملف الوورد النهائي
word_file_path = os.path.join(output_dir, "Normalization_Quality_Report.docx")

# إنشاء مستند وورد جديد
document = Document()
document.add_heading("Normalization Quality Evaluation Report", level=1)

# الحصول على جميع ملفات FIF في مجلد الإدخال
fif_files = glob.glob(os.path.join(input_dir, "*.fif"))

for file_path in fif_files:
    base_name = os.path.basename(file_path).replace(".fif", "")
    print("Processing file:", base_name)
    
    # إضافة عنوان فرعي باسم الملف في التقرير
    document.add_heading(base_name, level=2)
    
    # قراءة الملف
    raw = mne.io.read_raw_fif(file_path, preload=True, verbose=False)
    data, _ = raw[:]  # data.shape = (n_channels, n_samples)
    
    # حساب المتوسط والانحراف المعياري لكل قناة
    means = np.mean(data, axis=1)
    stds = np.std(data, axis=1)
    channel_names = raw.ch_names
    
    # إنشاء جدول في التقرير: الأعمدة (Channel, Mean, Std)
    table = document.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Channel"
    hdr_cells[1].text = "Mean"
    hdr_cells[2].text = "Std"
    
    for ch, m, s in zip(channel_names, means, stds):
        row_cells = table.add_row().cells
        row_cells[0].text = str(ch)
        row_cells[1].text = f"{m:.3f}"
        row_cells[2].text = f"{s:.3f}"
    
    # رسم الهيستوجرام للقيم (بشكل مشابه للمظهر المطلوب)
    flattened_data = data.flatten()
    plt.figure(figsize=(6, 4))
    plt.hist(flattened_data, bins=50, density=True, edgecolor='black')
    plt.title("Histogram of Normalized Data Values")
    plt.xlabel("Normalized Value")
    plt.ylabel("Density")
    
    # حفظ الشكل كصورة
    image_file = os.path.join(output_dir, f"{base_name}_histogram.png")
    plt.savefig(image_file, dpi=100, bbox_inches='tight')
    plt.close()
    
    # إدراج الصورة في مستند الوورد
    document.add_paragraph()  # فصل بين الجدول والصورة
    document.add_picture(image_file, width=Inches(6))
    document.add_paragraph()  # سطر فارغ
    document.add_page_break()  # فصل صفحة لكل ملف

# حفظ ملف الوورد
document.save(word_file_path)
print("Word report saved to:", word_file_path)
